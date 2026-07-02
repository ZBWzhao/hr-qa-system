import os
import re
import tempfile
from typing import Optional
from datetime import datetime
from fastapi import APIRouter, Depends, UploadFile, File, Form
from sqlalchemy.orm import Session
from app.core.database import get_db
from app.core.deps import get_current_user, require_roles
from app.core.response import success, error, paginated
from app.core.config import settings
from app.schemas.document import DocumentOut, ChunkOut
from app.models.document import Document, DocumentChunk, DocumentVersion
from app.models.user import User
from app.services.rag.vectorstore import add_documents, delete_document

router = APIRouter()


def split_text_to_chunks(text: str, chunk_size: int = 400) -> list:
    paragraphs = re.split(r'\n\s*\n', text)
    chunks = []
    current = ""
    for p in paragraphs:
        p = p.strip()
        if not p:
            continue
        if len(current) + len(p) > chunk_size and current:
            chunks.append(current.strip())
            current = p
        else:
            current = current + "\n\n" + p if current else p
    if current.strip():
        chunks.append(current.strip())
    final = []
    for chunk in chunks:
        if len(chunk) > chunk_size * 2:
            for i in range(0, len(chunk), chunk_size):
                final.append(chunk[i:i + chunk_size])
        else:
            final.append(chunk)
    return final


def extract_keywords(text: str) -> str:
    words = re.findall(r'[\u4e00-\u9fff]{2,6}', text)
    freq = {}
    for w in words:
        freq[w] = freq.get(w, 0) + 1
    sorted_words = sorted(freq.items(), key=lambda x: x[1], reverse=True)
    return ",".join([w for w, _ in sorted_words[:10]])


@router.get("")
def list_documents(keyword: Optional[str] = None, category: Optional[str] = None, status: Optional[str] = None, page: int = 1, page_size: int = 20, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    import re

    query = db.query(Document)
    # 普通员工只能看已发布文档
    if current_user.role not in ("hr", "admin"):
        query = query.filter(Document.status == "published")
    if keyword:
        # 同时搜索标题和内容
        query = query.filter(
            (Document.title.contains(keyword)) | (Document.content_text.contains(keyword))
        )
    if category:
        query = query.filter(Document.category == category)
    if status:
        query = query.filter(Document.status == status)
    total = query.count()
    items = query.order_by(Document.updated_at.desc()).offset((page - 1) * page_size).limit(page_size).all()

    def get_highlighted_content(content, keyword, context_len=150):
        """获取关键词周围的内容并高亮"""
        if not content or not keyword:
            return ""
        # 查找关键词位置
        idx = content.lower().find(keyword.lower())
        if idx == -1:
            return ""
        # 截取关键词周围的内容
        start = max(0, idx - context_len)
        end = min(len(content), idx + len(keyword) + context_len)
        snippet = content[start:end]
        if start > 0:
            snippet = "..." + snippet
        if end < len(content):
            snippet = snippet + "..."
        # 高亮关键词
        pattern = re.compile(re.escape(keyword), re.IGNORECASE)
        snippet = pattern.sub(f'<em>{keyword}</em>', snippet)
        return snippet

    result = []
    for d in items:
        doc_dict = DocumentOut.model_validate(d).model_dump()
        uploader = db.query(User).filter(User.id == d.uploader_id).first()
        doc_dict["uploader_name"] = uploader.real_name if uploader else "未知"
        # 如果有搜索关键词，添加高亮的内容片段
        if keyword:
            doc_dict["highlighted_content"] = get_highlighted_content(d.content_text, keyword)
        result.append(doc_dict)
    return paginated(result, total, page, page_size)


ALLOWED_EXTENSIONS = {"txt", "md", "docx", "pdf"}
BLOCKED_EXTENSIONS = {"exe", "zip", "bat", "cmd", "sh", "msi", "dll", "com", "scr"}


def extract_docx_text(file_path: str) -> str:
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)
    except Exception:
        return None


@router.post("/classify")
def classify_document(file: UploadFile = File(...)):
    """根据文档内容自动分类"""
    ext = file.filename.split(".")[-1].lower() if "." in file.filename else "unknown"
    content = ""

    try:
        if ext in ("txt", "md"):
            raw = file.file.read()
            for enc in ("utf-8", "utf-8-sig", "utf-16", "gbk", "gb2312", "latin-1"):
                try:
                    content = raw.decode(enc)
                    break
                except (UnicodeDecodeError, UnicodeError):
                    continue
        elif ext == "docx":
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                tmp.write(file.file.read())
                tmp_path = tmp.name
            content = extract_docx_text(tmp_path) or ""
            os.unlink(tmp_path)
        elif ext == "pdf":
            try:
                import fitz
                raw = file.file.read()
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                    tmp.write(raw)
                    tmp_path = tmp.name
                doc = fitz.open(tmp_path)
                content = "\n".join([page.get_text() for page in doc])
                doc.close()
                os.unlink(tmp_path)
            except ImportError:
                return error("PDF解析需要安装PyMuPDF库")
    except Exception as e:
        return error(f"文件读取失败: {str(e)}")

    if not content.strip():
        return success({"category": "other", "confidence": 0})

    # 分类关键词映射
    category_keywords = {
        "attendance": ["考勤", "打卡", "签到", "迟到", "早退", "旷工", "出勤", "加班", "工时", "排班"],
        "salary": ["薪酬", "工资", "薪资", "奖金", "绩效工资", "提成", "补贴", "津贴", "社保", "公积金", "五险一金"],
        "benefit": ["福利", "体检", "节日", "礼品", "团建", "活动", "保险", "商业保险", "补充医疗"],
        "leave": ["休假", "请假", "年假", "病假", "事假", "婚假", "产假", "陪产假", "丧假", "调休", "假期"],
        "performance": ["绩效", "考核", "KPI", "OKR", "目标", "评估", "评分", "晋升", "升职", "述职"]
    }

    scores = {}
    for cat, keywords in category_keywords.items():
        score = sum(content.count(kw) for kw in keywords)
        scores[cat] = score

    total = sum(scores.values())
    if total == 0:
        return success({"category": "other", "confidence": 0})

    best_cat = max(scores, key=scores.get)
    confidence = round(scores[best_cat] / total * 100)

    return success({"category": best_cat, "confidence": confidence})


@router.get("/{doc_id}")
def get_document(doc_id: int, db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        return error("文档不存在")
    doc_dict = DocumentOut.model_validate(doc).model_dump()
    uploader = db.query(User).filter(User.id == doc.uploader_id).first()
    doc_dict["uploader_name"] = uploader.real_name if uploader else "未知"
    versions = db.query(DocumentVersion).filter(DocumentVersion.document_id == doc_id).order_by(DocumentVersion.created_at.desc()).all()
    doc_dict["versions"] = [{"id": v.id, "version": v.version, "created_by": v.created_by, "created_at": v.created_at.isoformat()} for v in versions]
    return success(doc_dict)


@router.post("")
def create_document(title: str = Form(...), category: str = Form("other"), content_text: Optional[str] = Form(None), file: Optional[UploadFile] = File(None), current_user: User = Depends(require_roles("hr", "admin")), db: Session = Depends(get_db)):
    file_path = None
    file_type = None
    file_content = ""

    if file:
        ext = file.filename.split(".")[-1].lower() if "." in file.filename else "unknown"
        if ext in BLOCKED_EXTENSIONS:
            return error(f"不支持上传 {ext} 格式文件")
        if ext in ("xlsx", "xls"):
            return error(f"暂不支持解析 {ext} 格式，请使用 txt/md/docx/pdf 或手动录入")
        if ext not in ALLOWED_EXTENSIONS and ext != "pdf":
            return error(f"不支持上传 {ext} 格式文件")
        os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
        file_path = os.path.join(settings.UPLOAD_DIR, f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{file.filename}")
        with open(file_path, "wb") as f:
            f.write(file.file.read())
        file_type = ext
        if ext in ("txt", "md"):
            for enc in ("utf-8", "utf-8-sig", "utf-16", "gbk", "gb2312", "latin-1"):
                try:
                    file_content = open(file_path, "r", encoding=enc).read()
                    break
                except (UnicodeDecodeError, UnicodeError):
                    continue
            else:
                return error("文件编码无法识别，请使用 UTF-8 编码保存后重试")
        elif ext == "docx":
            file_content = extract_docx_text(file_path) or ""
        elif ext == "pdf":
            try:
                import fitz
                pdf_doc = fitz.open(file_path)
                file_content = "\n".join([page.get_text() for page in pdf_doc])
                pdf_doc.close()
            except ImportError:
                return error("PDF解析需要安装PyMuPDF库")
        if not title or title.strip() == "":
            title = file.filename.rsplit(".", 1)[0]

    # 合并文件内容和手动录入内容
    parts = []
    if file_content and file_content.strip():
        parts.append(file_content.strip())
    if content_text and content_text.strip():
        parts.append(content_text.strip())
    actual_content = "\n\n".join(parts) if parts else None

    if file and (not actual_content or not actual_content.strip()):
        return error("文件内容为空")

    doc = Document(
        title=title,
        category=category,
        file_path=file_path,
        file_type=file_type,
        content_text=actual_content,
        version="1.0",
        status="draft",
        uploader_id=current_user.id
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    if actual_content:
        chunks = split_text_to_chunks(actual_content)
        for i, chunk in enumerate(chunks):
            keywords = extract_keywords(chunk)
            doc_chunk = DocumentChunk(document_id=doc.id, chunk_index=i, content=chunk, keywords=keywords)
            db.add(doc_chunk)
        db.commit()

    version_record = DocumentVersion(document_id=doc.id, version="1.0", file_path=file_path, content_text=actual_content, created_by=current_user.id)
    db.add(version_record)
    db.commit()

    return success(DocumentOut.model_validate(doc).model_dump())


@router.put("/{doc_id}")
def update_document(doc_id: int, title: Optional[str] = Form(None), category: Optional[str] = Form(None), content_text: Optional[str] = Form(None), current_user: User = Depends(require_roles("hr", "admin")), db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        return error("文档不存在")
    # 保存旧版本
    old_version_record = DocumentVersion(
        document_id=doc.id, version=doc.version, file_path=doc.file_path,
        content_text=doc.content_text, created_by=current_user.id
    )
    db.add(old_version_record)
    if title is not None:
        doc.title = title
    if category is not None:
        doc.category = category
    if content_text is not None:
        doc.content_text = content_text
        db.query(DocumentChunk).filter(DocumentChunk.document_id == doc_id).delete()
        chunks = split_text_to_chunks(content_text)
        for i, chunk in enumerate(chunks):
            keywords = extract_keywords(chunk)
            doc_chunk = DocumentChunk(document_id=doc.id, chunk_index=i, content=chunk, keywords=keywords)
            db.add(doc_chunk)
    parts = doc.version.split(".")
    doc.version = f"{parts[0]}.{int(parts[1]) + 1}"
    # 保存新版本
    new_version_record = DocumentVersion(
        document_id=doc.id, version=doc.version, file_path=doc.file_path,
        content_text=doc.content_text, created_by=current_user.id
    )
    db.add(new_version_record)
    db.commit()
    db.refresh(doc)
    return success(DocumentOut.model_validate(doc).model_dump())


@router.delete("/{doc_id}")
def delete_document_endpoint(doc_id: int, current_user: User = Depends(require_roles("hr", "admin")), db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        return error("文档不存在")

    # 删除Chroma中的向量
    delete_document(doc_id)

    db.query(DocumentChunk).filter(DocumentChunk.document_id == doc_id).delete()
    db.query(DocumentVersion).filter(DocumentVersion.document_id == doc_id).delete()
    db.delete(doc)
    db.commit()
    return success(None, "删除成功")


@router.post("/{doc_id}/publish")
def publish_document(doc_id: int, current_user: User = Depends(require_roles("hr", "admin")), db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        return error("文档不存在")
    doc.status = "published"
    db.commit()

    # 发布时自动向量化到Chroma
    chunks = db.query(DocumentChunk).filter(DocumentChunk.document_id == doc_id).all()
    if chunks:
        chunk_dicts = [{"content": c.content, "keywords": c.keywords or ""} for c in chunks]
        add_documents(doc_id, chunk_dicts)

    return success(None, "发布成功")


@router.post("/{doc_id}/archive")
def archive_document(doc_id: int, current_user: User = Depends(require_roles("hr", "admin")), db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        return error("文档不存在")
    doc.status = "archived"
    db.commit()
    return success(None, "归档成功")


@router.post("/{doc_id}/unarchive")
def unarchive_document(doc_id: int, current_user: User = Depends(require_roles("hr", "admin")), db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        return error("文档不存在")
    doc.status = "draft"
    db.commit()
    return success(None, "下架成功")


@router.get("/{doc_id}/download")
def download_document(doc_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    """下载文档附件"""
    from fastapi.responses import FileResponse

    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        return error("文档不存在")
    if not doc.file_path:
        return error("该文档没有附件")

    import os
    if not os.path.exists(doc.file_path):
        return error("附件文件不存在，可能已被删除")

    # 获取文件名
    file_name = os.path.basename(doc.file_path)

    return FileResponse(
        path=doc.file_path,
        filename=file_name,
        media_type='application/octet-stream'
    )


@router.get("/{doc_id}/chunks")
def get_document_chunks(doc_id: int, db: Session = Depends(get_db)):
    chunks = db.query(DocumentChunk).filter(DocumentChunk.document_id == doc_id).order_by(DocumentChunk.chunk_index).all()
    return success([ChunkOut.model_validate(c).model_dump() for c in chunks])


@router.get("/{doc_id}/versions")
def get_document_versions(doc_id: int, db: Session = Depends(get_db)):
    versions = db.query(DocumentVersion).filter(DocumentVersion.document_id == doc_id).order_by(DocumentVersion.created_at.desc()).all()
    return success([{"id": v.id, "version": v.version, "created_by": v.created_by, "created_at": v.created_at.isoformat()} for v in versions])


@router.post("/reindex")
def reindex_all_documents(current_user: User = Depends(require_roles("hr", "admin")), db: Session = Depends(get_db)):
    """重新索引所有已发布文档到向量数据库"""
    from app.services.rag.vectorstore import get_collection_stats

    documents = db.query(Document).filter(Document.status == "published").all()
    indexed_count = 0

    for doc in documents:
        chunks = db.query(DocumentChunk).filter(DocumentChunk.document_id == doc.id).all()
        if chunks:
            chunk_dicts = [{"content": c.content, "keywords": c.keywords or ""} for c in chunks]
            add_documents(doc.id, chunk_dicts)
            indexed_count += 1

    stats = get_collection_stats()
    return success({
        "indexed_documents": indexed_count,
        "total_chunks": stats["total_chunks"]
    }, f"已重新索引 {indexed_count} 个文档")
